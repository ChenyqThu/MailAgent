"""anydoc lane（task 08-10 WP2）—— lane 判定 + 回落纪律 + flag-off 等价。

本文件钉的核心不变量：**关着的时候，这个功能必须像不存在一样**。
`extract_text()` 的每条原生分支一个字都没动，flag off 时产出必须与改动前逐字节相同；
开着但转换失败时也必须完整落回原生 extractor —— 最坏情况只能是「没变好」，不能是「变坏」。

另一条纪律：回落只认**异常类型与空产出**，绝不解析 anydoc 的错误字符串。
"""

from __future__ import annotations

from pathlib import Path

import pytest

from src.converter import anydoc_extract, attachment_text
from src.converter.anydoc_extract import (
    ANYDOC_LEGACY_EXTENSIONS,
    ANYDOC_OFFICE_EXTENSIONS,
    LANE_LEGACY,
    LANE_OFFICE,
    LANE_PDF,
    lane_for_extension,
)
from src.converter.attachment_text import extract_text


def _make_docx(path: Path) -> Path:
    import docx as python_docx

    doc = python_docx.Document()
    doc.add_paragraph("一段中文")
    doc.add_paragraph("English paragraph")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Col A"
    tbl.cell(0, 1).text = "Col B"
    tbl.cell(1, 0).text = "v1"
    tbl.cell(1, 1).text = "v2"
    doc.save(str(path))
    return path


def _enable(monkeypatch, *, lanes: str = "office,legacy") -> None:
    monkeypatch.setattr(anydoc_extract, 'anydoc_enabled', lambda: True)
    monkeypatch.setattr(
        anydoc_extract, 'enabled_lanes',
        lambda: frozenset(x.strip() for x in lanes.split(',') if x.strip()),
    )


class TestLaneMapping:
    def test_office_and_legacy_and_pdf(self):
        assert lane_for_extension('.docx') == LANE_OFFICE
        assert lane_for_extension('.xlsm') == LANE_OFFICE
        assert lane_for_extension('.doc') == LANE_LEGACY
        assert lane_for_extension('.pdf') == LANE_PDF

    def test_case_insensitive(self):
        assert lane_for_extension('.DOCX') == LANE_OFFICE

    def test_csv_deliberately_unmapped(self):
        """🔴 anydoc 支持 csv，但现状 plaintext 直读已是最忠实产出，换过去只增风险。"""
        assert lane_for_extension('.csv') is None

    def test_images_and_plaintext_unmapped(self):
        for ext in ('.png', '.jpg', '.txt', '.md', '.log', '.zip'):
            assert lane_for_extension(ext) is None, ext

    def test_lane_sets_disjoint(self):
        assert not (ANYDOC_OFFICE_EXTENSIONS & ANYDOC_LEGACY_EXTENSIONS)


class TestLanesConfigParsing:
    def test_unknown_lane_ignored_not_fatal(self, monkeypatch):
        """应急阀写错一个字不该让整条提取链起不来。"""
        monkeypatch.setattr(
            anydoc_extract, '_config',
            lambda: type('C', (), {'mailagent_anydoc_lanes': 'office, bogus ,pdf'})(),
        )
        assert anydoc_extract.enabled_lanes() == frozenset({LANE_OFFICE, LANE_PDF})

    def test_empty_means_no_lane(self, monkeypatch):
        monkeypatch.setattr(
            anydoc_extract, '_config',
            lambda: type('C', (), {'mailagent_anydoc_lanes': ''})(),
        )
        assert anydoc_extract.enabled_lanes() == frozenset()

    def test_pdf_not_in_shipped_default(self):
        """🔴 实测 25 份真实 PDF：3 份回归，其中伪粗体重复病无判据可拦。默认不带 pdf。"""
        from src.config import Config

        default = Config.model_fields['mailagent_anydoc_lanes'].default
        assert LANE_PDF not in default
        assert LANE_OFFICE in default and LANE_LEGACY in default

    def test_flag_ships_off(self):
        from src.config import Config

        assert Config.model_fields['mailagent_anydoc_enabled'].default is False


class TestFlagOffEquivalence:
    """PRD 验收标准 1：flag off ⇒ 与改动前逐字节相同。"""

    def test_docx_uses_native_extractor(self, tmp_path: Path, monkeypatch):
        monkeypatch.setattr(anydoc_extract, 'anydoc_enabled', lambda: False)
        r = extract_text(_make_docx(tmp_path / "demo.docx"))
        assert r.extractor == 'docx'
        assert r.status == 'extracted'
        assert '一段中文' in r.text

    def test_anydoc_never_invoked_when_off(self, tmp_path: Path, monkeypatch):
        monkeypatch.setattr(anydoc_extract, 'anydoc_enabled', lambda: False)

        def _boom(*a, **k):  # pragma: no cover - 被调用即测试失败
            raise AssertionError('anydoc must not be touched while the flag is off')

        monkeypatch.setattr(anydoc_extract, 'convert_path', _boom)
        assert extract_text(_make_docx(tmp_path / "demo.docx")).extractor == 'docx'

    def test_lane_off_for_pdf_by_default(self, tmp_path: Path, monkeypatch):
        """总开关开着、但 pdf 不在 lanes 里 ⇒ PDF 仍归 pypdf。"""
        _enable(monkeypatch, lanes='office,legacy')

        def _boom(*a, **k):  # pragma: no cover
            raise AssertionError('pdf lane is disabled; anydoc must not run')

        monkeypatch.setattr(anydoc_extract, 'convert_path', _boom)
        f = tmp_path / "x.pdf"
        f.write_bytes(b"%PDF-1.4 not a real pdf")
        r = extract_text(f)
        assert r.extractor in ('pypdf', 'pdf_ocr', 'none')


class TestAnydocLaneActive:
    def test_docx_goes_through_anydoc(self, tmp_path: Path, monkeypatch):
        _enable(monkeypatch)
        monkeypatch.setattr(
            anydoc_extract, 'convert_path',
            lambda p: '# Title\n\n| a | b |\n| --- | --- |\n| 1 | 2 |',
        )
        r = extract_text(_make_docx(tmp_path / "demo.docx"))
        assert r.extractor == 'anydoc'
        assert r.status == 'extracted'
        assert '| --- |' in r.text

    def test_legacy_doc_no_longer_needs_soffice(self, tmp_path: Path, monkeypatch):
        """老 .doc 走 anydoc ⇒ 不再触碰没打包进 .app 的系统级 LibreOffice。"""
        _enable(monkeypatch)
        monkeypatch.setattr(anydoc_extract, 'convert_path', lambda p: 'legacy content')

        from src.converter import office_converter

        def _boom(*a, **k):  # pragma: no cover
            raise AssertionError('soffice bridge must not run when the legacy lane handles it')

        monkeypatch.setattr(office_converter, '_run_soffice_convert', _boom)
        f = tmp_path / "old.doc"
        f.write_bytes(b"\xd0\xcf\x11\xe0fake ole")
        r = extract_text(f)
        assert r.extractor == 'anydoc'
        assert r.text == 'legacy content'

    def test_truncation_applies(self, tmp_path: Path, monkeypatch):
        _enable(monkeypatch)
        huge = 'x' * (attachment_text.ATTACHMENT_TEXT_MAX_BYTES + 5000)
        monkeypatch.setattr(anydoc_extract, 'convert_path', lambda p: huge)
        r = extract_text(_make_docx(tmp_path / "demo.docx"))
        assert r.extractor == 'anydoc'
        assert r.truncated is True
        assert len(r.text.encode('utf-8')) <= attachment_text.ATTACHMENT_TEXT_MAX_BYTES


class TestFallback:
    """anydoc 失败必须完整落回原生 extractor —— 最坏是「没变好」，不能是「变坏」。"""

    def test_convert_failure_falls_back_to_native(self, tmp_path: Path, monkeypatch):
        _enable(monkeypatch)
        monkeypatch.setattr(anydoc_extract, 'convert_path', lambda p: None)
        r = extract_text(_make_docx(tmp_path / "demo.docx"))
        assert r.extractor == 'docx'
        assert '一段中文' in r.text

    def test_empty_output_falls_back(self, tmp_path: Path, monkeypatch):
        _enable(monkeypatch)
        monkeypatch.setattr(anydoc_extract, 'convert_path', lambda p: '   \n  ')
        r = extract_text(_make_docx(tmp_path / "demo.docx"))
        assert r.extractor == 'docx'

    def test_convert_path_swallows_exceptions(self, tmp_path: Path, monkeypatch):
        """回落信号是异常类型本身，不解析错误字符串；异常绝不越过本层。"""
        fake = type('M', (), {
            'to_markdown': staticmethod(
                lambda p: (_ for _ in ()).throw(RuntimeError('scanned, 17 pages'))
            )
        })
        monkeypatch.setattr(anydoc_extract, '_load', lambda: fake)
        assert anydoc_extract.convert_path(tmp_path / 'whatever.docx') is None

    def test_missing_package_soft_lands(self, tmp_path: Path, monkeypatch):
        _enable(monkeypatch)
        monkeypatch.setattr(anydoc_extract, '_load', lambda: None)
        r = extract_text(_make_docx(tmp_path / "demo.docx"))
        assert r.extractor == 'docx'
        assert anydoc_extract.available() is False


class TestRealAnydoc:
    """真跑一次装好的 anydoc —— mock 测不出「上游 API 换了签名」。"""

    def test_real_docx_produces_valid_gfm_table(self, tmp_path: Path, monkeypatch):
        pytest.importorskip('anydoc')
        _enable(monkeypatch)
        r = extract_text(_make_docx(tmp_path / "demo.docx"))
        assert r.extractor == 'anydoc'
        assert r.status == 'extracted'
        assert '一段中文' in r.text
        # 现状 python-docx 拼的是 'Col A | Col B' —— 缺 |---| 分隔行，不是合法 GFM。
        assert '| --- |' in r.text or '|---|' in r.text

    def test_bytes_api_matches_path_api(self, tmp_path: Path):
        """WP3 靠 to_markdown_bytes 免落盘；两条路产出必须一致，否则会出最难查的缝。"""
        pytest.importorskip('anydoc')
        f = _make_docx(tmp_path / "demo.docx")
        assert anydoc_extract.convert_bytes(f.read_bytes()) == anydoc_extract.convert_path(f)

    def test_garbage_bytes_return_none_not_raise(self):
        pytest.importorskip('anydoc')
        assert anydoc_extract.convert_bytes(b'\x00\x01not a document') is None
        assert anydoc_extract.convert_bytes(b'') is None
