"""POST /api/attachment/convert — chat 附件的内存文档 → markdown（task 08-10 WP3）。

这个端点与本 router 其余端点根本不同：**输入不是已入库的附件**，没有 attachment_id、
不查库、不碰 AttachmentStore。就是「一段字节进、一段 markdown 出」。

本文件钉三件事：
  1. flag 关着 ⇒ 404（renderer 据此静默回落 metadata-only，不给用户看半截功能）；
  2. 转换失败/格式不支持 **不是 4xx** ⇒ 200 + status='unsupported'，
     这样用户的消息照常发得出去，只是模型看不到附件正文；
  3. 大小护栏在**解码之前**先按 base64 长度估算 —— 把防 OOM 的闸放在会 OOM 的
     操作之后就等于没有。
"""

from __future__ import annotations

import base64
import io

import pytest

from src.converter import anydoc_extract


def _docx_bytes() -> bytes:
    import docx as python_docx

    doc = python_docx.Document()
    doc.add_paragraph("合同正文第一段")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "项目"
    tbl.cell(0, 1).text = "金额"
    tbl.cell(1, 0).text = "服务费"
    tbl.cell(1, 1).text = "1000"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _post(client, payload):
    return client.post("/api/attachment/convert", json=payload)


@pytest.fixture()
def anydoc_on(monkeypatch):
    monkeypatch.setattr(anydoc_extract, 'anydoc_enabled', lambda: True)
    monkeypatch.setattr(
        anydoc_extract, 'enabled_lanes',
        lambda: frozenset({anydoc_extract.LANE_OFFICE, anydoc_extract.LANE_LEGACY}),
    )


class TestFlagGate:
    def test_disabled_returns_404(self, client, monkeypatch):
        monkeypatch.setattr(anydoc_extract, 'anydoc_enabled', lambda: False)
        r = _post(client, {"filename": "a.docx", "contentBase64": "AAAA"})
        assert r.status_code == 404
        assert r.json()["error"]["code"] == "E_NOT_FOUND"

    def test_disabled_checked_before_payload_validation(self, client, monkeypatch):
        """关着的时候连参数都不该看 —— 否则会从错误码里泄漏端点存在与否的差异。"""
        monkeypatch.setattr(anydoc_extract, 'anydoc_enabled', lambda: False)
        assert _post(client, {}).status_code == 404


class TestValidation:
    def test_missing_filename(self, client, anydoc_on):
        r = _post(client, {"contentBase64": "AAAA"})
        assert r.status_code == 400
        assert r.json()["error"]["code"] == "E_INVALID_ARG"

    def test_missing_content(self, client, anydoc_on):
        r = _post(client, {"filename": "a.docx"})
        assert r.status_code == 400

    def test_invalid_base64(self, client, anydoc_on):
        r = _post(client, {"filename": "a.docx", "contentBase64": "not!valid!base64!"})
        assert r.status_code == 400
        assert r.json()["error"]["code"] == "E_INVALID_ARG"

    def test_oversized_rejected_before_decode(self, client, anydoc_on):
        """20 MiB 的 base64 串必须在 b64decode 之前就被拒 —— 护栏晚一步就等于没有。"""
        huge = "A" * (21 * 1024 * 1024)
        r = _post(client, {"filename": "a.docx", "contentBase64": huge})
        assert r.status_code == 413
        assert r.json()["error"]["code"] == "E_TOO_LARGE"


class TestConversion:
    def test_docx_converts_to_gfm(self, client, anydoc_on):
        pytest.importorskip('anydoc')
        payload = {
            "filename": "contract.docx",
            "contentBase64": base64.b64encode(_docx_bytes()).decode(),
        }
        r = _post(client, payload)
        assert r.status_code == 200
        data = r.json()["data"]
        assert data["status"] == "converted"
        assert data["extractor"] == "anydoc"
        assert "合同正文第一段" in data["markdown"]
        # 合法 GFM 表格（现状 python-docx 拼的 `a | b` 缺这一行）
        assert "| --- |" in data["markdown"] or "|---|" in data["markdown"]

    def test_unsupported_extension_is_200_not_error(self, client, anydoc_on):
        """回落必须是「没内容」而不是「报错」，否则用户的消息发不出去。"""
        r = _post(client, {
            "filename": "photo.png",
            "contentBase64": base64.b64encode(b"\x89PNG\r\n\x1a\n").decode(),
        })
        assert r.status_code == 200
        data = r.json()["data"]
        assert data["status"] == "unsupported"
        assert data["markdown"] is None

    def test_pdf_unsupported_while_lane_off(self, client, anydoc_on):
        """pdf lane 默认不开 ⇒ 即便总开关开着，PDF 也走不进转换。"""
        r = _post(client, {
            "filename": "doc.pdf",
            "contentBase64": base64.b64encode(b"%PDF-1.4 whatever").decode(),
        })
        assert r.json()["data"]["status"] == "unsupported"

    def test_corrupt_document_is_200_unsupported(self, client, anydoc_on):
        pytest.importorskip('anydoc')
        r = _post(client, {
            "filename": "broken.docx",
            "contentBase64": base64.b64encode(b"PK\x03\x04truncated garbage").decode(),
        })
        assert r.status_code == 200
        assert r.json()["data"]["status"] == "unsupported"

    def test_long_output_truncated_and_flagged(self, client, anydoc_on, monkeypatch):
        from src.api.routers.attachment import CHAT_ATTACHMENT_MAX_CHARS

        monkeypatch.setattr(
            anydoc_extract, 'convert_bytes',
            lambda data, fmt=None: 'x' * (CHAT_ATTACHMENT_MAX_CHARS + 500),
        )
        r = _post(client, {
            "filename": "big.docx", "contentBase64": base64.b64encode(b"anything").decode(),
        })
        data = r.json()["data"]
        assert data["status"] == "converted"
        assert data["truncated"] is True
        assert len(data["markdown"]) == CHAT_ATTACHMENT_MAX_CHARS

    def test_nothing_written_to_disk(self, client, anydoc_on, tmp_path, monkeypatch):
        """🔴 chat 附件从未落盘是本端点的安全前提 —— 转换不得留下任何文件。"""
        pytest.importorskip('anydoc')
        monkeypatch.chdir(tmp_path)
        before = set(tmp_path.rglob('*'))
        _post(client, {
            "filename": "contract.docx",
            "contentBase64": base64.b64encode(_docx_bytes()).decode(),
        })
        assert set(tmp_path.rglob('*')) == before
