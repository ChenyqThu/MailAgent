#!/usr/bin/env python3
"""测试 Office 文档转换功能"""

import sys
import os
import tempfile
from pathlib import Path

# 添加项目根目录到 path
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.converter.office_converter import (
    is_convertible,
    convert_to_csv,
    convert_to_pdf,
    check_soffice_available,
)


def test_is_convertible():
    """测试文件类型识别"""
    assert is_convertible("report.docx") == True
    assert is_convertible("slides.pptx") == True
    assert is_convertible("data.xlsx") == True
    assert is_convertible("image.png") == False
    assert is_convertible("doc.pdf") == False
    assert is_convertible("REPORT.DOCX") == True  # 大小写
    print("[PASS] is_convertible")


def test_xlsx_to_csv():
    """测试 xlsx → csv 转换"""
    try:
        import pandas as pd
    except ImportError:
        print("[SKIP] pandas not installed")
        return

    # 创建测试 xlsx
    with tempfile.TemporaryDirectory() as tmpdir:
        xlsx_path = os.path.join(tmpdir, "test_data.xlsx")

        # 创建多 sheet 测试文件
        with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
            df1 = pd.DataFrame({
                "姓名": ["张三", "李四", "王五"],
                "Score": [95, 87, 92],
                "Email": ["a@test.com", "b@test.com", "c@test.com"],
            })
            df1.to_excel(writer, sheet_name="成绩单", index=False)

            df2 = pd.DataFrame({
                "Item": ["A", "B"],
                "Value": [100, 200],
            })
            df2.to_excel(writer, sheet_name="Summary", index=False)

        print(f"Created test xlsx: {xlsx_path} ({Path(xlsx_path).stat().st_size} bytes)")

        # 转换
        output_dir = os.path.join(tmpdir, "output")
        os.makedirs(output_dir)
        csv_paths = convert_to_csv(xlsx_path, output_dir)

        assert len(csv_paths) == 2, f"Expected 2 CSV files, got {len(csv_paths)}"

        for csv_path in csv_paths:
            p = Path(csv_path)
            assert p.exists(), f"CSV file not found: {csv_path}"
            content = p.read_text(encoding="utf-8-sig")
            print(f"  {p.name} ({p.stat().st_size} bytes):")
            for line in content.strip().split("\n")[:3]:
                print(f"    {line}")

        # 验证 CJK 内容
        first_csv = Path(csv_paths[0]).read_text(encoding="utf-8-sig")
        assert "张三" in first_csv, "CJK content not found in CSV"

        print("[PASS] xlsx_to_csv (multi-sheet, CJK)")


def test_docx_to_pdf():
    """测试 docx → pdf 转换（需要 LibreOffice）"""
    if not check_soffice_available():
        print("[SKIP] LibreOffice not available")
        return

    try:
        from docx import Document
    except ImportError:
        # 用 python-docx 创建测试文件；没装的话用简单 zip 模拟
        print("[SKIP] python-docx not installed, creating minimal docx manually")

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "test_doc.docx")

        try:
            from docx import Document
            doc = Document()
            doc.add_heading("测试文档", level=1)
            doc.add_paragraph("这是一个包含中文的测试文档。")
            doc.add_paragraph("Hello World - English text.")
            doc.save(docx_path)
        except ImportError:
            # 创建最简 docx（实际是 zip with XML）
            import zipfile
            with zipfile.ZipFile(docx_path, 'w') as zf:
                zf.writestr('[Content_Types].xml',
                    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                    '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                    '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                    '<Default Extension="xml" ContentType="application/xml"/>'
                    '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                    '</Types>')
                zf.writestr('_rels/.rels',
                    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                    '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                    '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
                    '</Relationships>')
                zf.writestr('word/document.xml',
                    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                    '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    '<w:body><w:p><w:r><w:t>测试文档 Test Document</w:t></w:r></w:p></w:body>'
                    '</w:document>')

        print(f"Created test docx: {docx_path} ({Path(docx_path).stat().st_size} bytes)")

        output_dir = os.path.join(tmpdir, "output")
        os.makedirs(output_dir)
        pdf_path = convert_to_pdf(docx_path, output_dir)

        if pdf_path:
            p = Path(pdf_path)
            assert p.exists(), f"PDF file not found: {pdf_path}"
            assert p.stat().st_size > 0, "PDF file is empty"
            # 检查 PDF magic bytes
            with open(pdf_path, 'rb') as f:
                header = f.read(5)
            assert header == b'%PDF-', f"Not a valid PDF: {header}"
            print(f"  {p.name} ({p.stat().st_size} bytes) - valid PDF")
            print("[PASS] docx_to_pdf")
        else:
            print("[FAIL] docx_to_pdf - conversion returned None")


if __name__ == "__main__":
    print("=" * 50)
    print("Office Converter Tests")
    print("=" * 50)

    test_is_convertible()
    test_xlsx_to_csv()
    test_docx_to_pdf()

    print("\n" + "=" * 50)
    print("All tests completed!")
